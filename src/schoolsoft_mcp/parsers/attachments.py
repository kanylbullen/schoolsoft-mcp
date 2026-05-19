"""Attachment download + text extraction.

SchoolSoft serves news/message attachments through a two-step flow:

1. ``jsp/student/right_student_file_download.jsp?requestid1=<news_id>
   &requestid2=<type_id>&object=<news|message>&fileid=<file_id>``
   returns a 302 to
2. ``/files/<school>/tmp_file_<file_id>.tmp?md5=<sig>&expires=<ts>&filename=<name>``
   which is the actual file. The signed URL expires after a few minutes.

We hide that two-step dance behind ``download_attachment_bytes`` —
``SchoolSoftClient.fetch_bytes`` follows the redirect transparently.
"""

from __future__ import annotations

import io
import logging
import mimetypes
import re
from typing import Any
from urllib.parse import unquote

logger = logging.getLogger(__name__)

# Eager-import the heavy text-extraction libs at module load instead of
# inside the per-call helpers. First call would otherwise pay the import
# cost (plus Windows AV file-scan on the freshly-installed .pyd files),
# which has been observed to take minutes on first use — long enough for
# Claude Desktop to cancel the tool call. Once the server process is up,
# subsequent calls reuse the already-imported modules.
try:
    from pypdf import PdfReader as _PdfReader
except ImportError:  # pragma: no cover - optional dep
    _PdfReader = None  # type: ignore[assignment,misc]

try:
    from docx import Document as _DocxDocument
except ImportError:  # pragma: no cover - optional dep
    _DocxDocument = None  # type: ignore[assignment]

# Hard cap on text extraction output. LLM contexts are finite; PDFs of
# hundreds of pages would be useless and expensive. Configurable per-call.
DEFAULT_TEXT_LIMIT = 50_000

ObjectKind = str  # "news" | "message"


def build_download_path(
    *,
    parent_id: int,
    type_id: int,
    fileid: int,
    object_kind: ObjectKind = "news",
) -> tuple[str, dict[str, str]]:
    """Return (path, params) for SchoolSoftClient.fetch_bytes."""
    return (
        "jsp/student/right_student_file_download.jsp",
        {
            "requestid1": str(parent_id),
            "requestid2": str(type_id),
            "object": object_kind,
            "fileid": str(fileid),
        },
    )


def guess_content_type(filename: str) -> str:
    """Best-effort content-type from filename. Falls back to octet-stream."""
    ctype, _ = mimetypes.guess_type(filename)
    return ctype or "application/octet-stream"


def filename_from_headers(headers: dict[str, str], fallback: str) -> str:
    """Extract a clean filename from a Content-Disposition header.

    Returns ``fallback`` when the header is missing or unparseable.
    Callers that want a URL-based fallback (e.g. the ``filename=`` query
    param on SchoolSoft's signed download URLs) should pass that as
    ``fallback`` themselves.
    """
    cd = headers.get("content-disposition", "")
    if not cd:
        return fallback
    # filename*=UTF-8''url-encoded wins over filename=...
    star = re.search(r"filename\*\s*=\s*[^']*''([^;]+)", cd, re.IGNORECASE)
    if star:
        return unquote(star.group(1).strip().strip('"'))
    plain = re.search(r'filename\s*=\s*"?([^";]+)"?', cd, re.IGNORECASE)
    if plain:
        return plain.group(1).strip()
    return fallback


def extract_text(
    content: bytes,
    content_type: str,
    *,
    limit: int = DEFAULT_TEXT_LIMIT,
) -> tuple[str, bool, str | None]:
    """Best-effort plain-text extraction. Returns (text, truncated, note).

    Supports PDF (via pypdf), .docx (via python-docx), and plain text. For
    other content types the note explains what to do (typically: use
    download_attachment to get the raw bytes).

    Routing is primarily by magic bytes — Content-Type is unreliable in
    practice (SchoolSoft's redirect chain produces values like
    ``"application/octet-stream, application/pdf"``).
    """
    # Magic bytes are authoritative for the binary formats we care about.
    if content[:4] == b"%PDF":
        return _truncate(_extract_pdf(content), limit)
    if content[:4] == b"PK\x03\x04" and _looks_like_docx(content):
        return _truncate(_extract_docx(content), limit)

    # Parse all comma- and semicolon-separated content-type variants and
    # accept the first plausible one.
    ctypes = _parse_content_types(content_type)
    if any(c == "application/pdf" for c in ctypes):
        return _truncate(_extract_pdf(content), limit)
    if any(
        c == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        for c in ctypes
    ):
        return _truncate(_extract_docx(content), limit)
    if any(c.startswith("text/") or c in {"application/json", "application/xml"} for c in ctypes):
        try:
            return _truncate(content.decode("utf-8"), limit)
        except UnicodeDecodeError:
            return _truncate(content.decode("latin-1", errors="replace"), limit)

    note = (
        f"No text extractor for content-type {content_type!r} "
        f"(first {len(content[:4])} bytes: {content[:4]!r}). "
        "Use download_attachment to fetch the raw bytes."
    )
    return "", False, note


def _parse_content_types(header: str) -> list[str]:
    """Split a Content-Type header into its constituent MIME types.

    Tolerates the comma-joined values SchoolSoft's redirect chain produces
    (e.g. ``"application/octet-stream, application/pdf"``) as well as the
    normal ``"text/html; charset=utf-8"`` form.
    """
    out: list[str] = []
    for chunk in header.split(","):
        primary = chunk.split(";", 1)[0].strip().lower()
        if primary:
            out.append(primary)
    return out


def _looks_like_docx(content: bytes) -> bool:
    """Detect a .docx by peeking for ``word/`` inside the ZIP central directory.

    A bare ZIP magic (``PK\\x03\\x04``) could be any Office Open XML file
    (.xlsx, .pptx) or a plain .zip. Look for the ``word/`` entry name in
    the first ~32 KB to be sure we hand it to python-docx.
    """
    head = content[: min(len(content), 32_768)]
    return b"word/" in head


def _truncate(text: str, limit: int) -> tuple[str, bool, str | None]:
    if len(text) <= limit:
        return text, False, None
    return (
        text[:limit] + f"\n\n... [truncated, {len(text) - limit} chars omitted]",
        True,
        None,
    )


def _extract_pdf(content: bytes) -> str:
    if _PdfReader is None:
        return (
            "[pypdf not installed; reinstall the package to enable PDF text "
            "extraction: pip install -e .]"
        )

    try:
        reader = _PdfReader(io.BytesIO(content))
    except Exception as err:
        logger.warning("PDF parse failed: %s", err)
        return f"[Could not parse PDF: {err}]"

    parts: list[str] = []
    for page_num, page in enumerate(reader.pages, start=1):
        try:
            parts.append(page.extract_text() or "")
        except Exception as err:
            logger.warning("PDF page %d extract failed: %s", page_num, err)
            parts.append(f"[page {page_num}: extraction failed]")
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(content: bytes) -> str:
    if _DocxDocument is None:
        return (
            "[python-docx not installed; reinstall the package to enable "
            ".docx text extraction: pip install -e .]"
        )

    try:
        doc: Any = _DocxDocument(io.BytesIO(content))
    except Exception as err:
        logger.warning("DOCX parse failed: %s", err)
        return f"[Could not parse DOCX: {err}]"

    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)
